import docx
import os
import shutil

doc_path = r"c:\Users\Topo_4\Documents\AT_PFE\Plan Détaillé du Mémoire de PFE.docx"
backup_path = r"c:\Users\Topo_4\Documents\AT_PFE\Plan Détaillé du Mémoire de PFE_backup.docx"

if os.path.exists(doc_path) and not os.path.exists(backup_path):
    shutil.copy2(doc_path, backup_path)

doc = docx.Document()

doc.add_heading("Plan provisoire détaillé du Mémoire de PFE", level=0)

doc.add_paragraph("Titre : Développement d'un outil permettant le traitement et l'insertion des archives numériques sur Geofoncier au sein d'un cabinet de Géomètre-Expert")

doc.add_heading("I. INTRODUCTION", level=1)
doc.add_heading("1.1. Contexte du projet", level=2)
doc.add_paragraph("L’importance des archives dans le métier de Géomètre-Expert (foncier, bornage, historique).", style='List Bullet')
doc.add_paragraph("La transition numérique des cabinets.", style='List Bullet')

doc.add_heading("1.2. La problématique", level=2)
doc.add_paragraph("La difficulté d'exploitation des registres anciens (manuscrits) et la diversité structurelle des plans modernes numérisés.", style='List Bullet')
doc.add_paragraph("Le coût humain et le risque d'erreur de la saisie manuelle pour Geofoncier.", style='List Bullet')

doc.add_heading("1.3. Objectifs et livrables", level=2)
doc.add_paragraph("Créer un processus hybride et autonome de détection, de lecture sémantique et de validation interactive.", style='List Bullet')

doc.add_heading("II. ANALYSE DU MÉTIER ET DES DONNÉES", level=1)
doc.add_heading("2.1. Le cadre de Geofoncier", level=2)
doc.add_paragraph("Rôle et fonctionnement du portail, intérêt du versement des données pour les confrères.", style='List Bullet')
doc.add_paragraph("Spécifications techniques des données attendues par l'API de l'Ordre (format, tri de pertinence…).", style='List Bullet')

doc.add_heading("2.2. Typologie des archives à traiter", level=2)
doc.add_paragraph("Analyse de la bivalence des flux documentaires :", style='List Bullet')
p = doc.add_paragraph("Les registres et livrets anciens (format manuscrit, papier, encres).")
# fallback for python-docx style
try:
    p.style = 'List Bullet 2'
except:
    pass
p = doc.add_paragraph("Les plans cadastraux modernes (documents DGFIP, DMPC, PVa).")
try:
    p.style = 'List Bullet 2'
except:
    pass
doc.add_paragraph("Analyse de la complexité des écritures et des variations de mise en page.", style='List Bullet')

doc.add_heading("III. ÉTAT DE L'ART", level=1)
doc.add_heading("3.1. Vision par ordinateur", level=2)
doc.add_paragraph("Détection d’objets : Pourquoi YOLO pour segmenter la structure des documents.", style='List Bullet')

doc.add_heading("3.2. Reconnaissance de l'écriture manuscrite (HTR)", level=2)
doc.add_paragraph("Les réseaux de neurones récurrents vs les Transformers.", style='List Bullet')
doc.add_paragraph("Étude comparative : TrOCR vs Kraken vs EasyOCR (en open source).", style='List Bullet')

doc.add_heading("3.3. L'avènement des modèles Vision-Langage (VLM)", level=2)
doc.add_paragraph("L'apport décisif des modèles multimodaux (ex: Qwen2-VL, LLaVA) pour la compréhension sémantique du texte dans son contexte visuel.", style='List Bullet')

doc.add_heading("IV. ARCHITECTURE DU PROJET", level=1)
doc.add_heading("4.1. Étape 1 : Classification et routage des documents", level=2)
doc.add_paragraph("Identification automatique de la nature de l'archive (plan moderne vs registre ancien) pour orienter le pipeline de traitement.", style='List Bullet')

doc.add_heading("4.2. Étape 2 : Extraction géométrique et segmentation", level=2)
doc.add_paragraph("Détection des colonnes et cellules via YOLO pour les livrets anciens (algorithme \"Ghost Lines\" pour reconstruire la structure).", style='List Bullet')
doc.add_paragraph("Recherche de correspondances géométriques et par ancrage de mots-clés pour le recadrage des plans modernes.", style='List Bullet')

doc.add_heading("4.3. Étape 3 : Moteur d'extraction hybride (HTR & VLM)", level=2)
doc.add_paragraph("Stratégie multi-hypothèses : Génération de plusieurs prédictions par TrOCR pour les écritures très dégradées.", style='List Bullet')
doc.add_paragraph("Utilisation des modèles Vision-Langage (VLM) en tant que moteur d'extraction sémantique direct pour les métadonnées complexes (N° DA, Section).", style='List Bullet')

doc.add_heading("4.4. Protection de données sensibles et dispositions mises en place", level=2)

doc.add_heading("V. FIABILISATION ET POST-TRAITEMENT DES DONNÉES", level=1)
doc.add_heading("5.1. Décodage et correction à la volée", level=2)
doc.add_paragraph("Utilisation des LogitsProcessors pour interdire les prédictions hors-dictionnaire ou hors communal.", style='List Bullet')
doc.add_paragraph("Normalisation syntaxique des toponymes (gestion des accents, majuscules).", style='List Bullet')

doc.add_heading("5.2. Matching intelligent et croisement d'archives", level=2)
doc.add_paragraph("Matching Fuzzy et similarité visuelle : pondération de la distance de Levenshtein par les confusions de lettres courantes.", style='List Bullet')
doc.add_paragraph("Croisement avec des répertoires historiques externes (ex: index des archives Racat & Ceyte) pour consolider la fiabilité des extractions.", style='List Bullet')

doc.add_heading("5.3. Interface \"Human-in-the-Loop\" : l'application de validation", level=2)
doc.add_paragraph("Développement d'une interface de contrôle réactive (Streamlit) garantissant l'intégrité de la donnée.", style='List Bullet')
doc.add_paragraph("Validation des métadonnées via les référentiels officiels (INSEE, liste nationale des géomètres) et auto-sauvegarde des corrections en temps réel.", style='List Bullet')

doc.add_heading("VI. INTÉGRATION SUR GEOFONCIER ET RÉSULTATS", level=1)
doc.add_heading("6.1. Le connecteur API Geofoncier", level=2)
doc.add_paragraph("Développement d'une communication directe avec les serveurs de Geofoncier (mise en place d'un mode Dry Run de simulation, requêtes multipart).", style='List Bullet')
doc.add_paragraph("Structuration, contrôle et formatage automatiques des paquets de données (JSON/CSV) pour l'API.", style='List Bullet')

doc.add_heading("6.2. Évaluation des performances", level=2)
doc.add_paragraph("Définition des métriques : Taux de succès par commune etc ...", style='List Bullet')
doc.add_paragraph("Comparaison de productivité : Gain de temps réel estimé pour le cabinet.", style='List Bullet')

doc.add_heading("6.3. Analyse des limites", level=2)
doc.add_paragraph("Gestion des faux positifs et nécessité du contrôle humain final, et à quel degré.", style='List Bullet')

doc.add_heading("VII. CONCLUSION ET PERSPECTIVES", level=1)
doc.add_paragraph("Synthèse des contributions techniques.", style='List Bullet')
doc.add_paragraph("Apports personnels du projet (IA, développement, métier de géomètre).", style='List Bullet')
doc.add_paragraph("Ouvertures : automatisation de la lecture des limites de propriétés sur les plans parcellaires ?", style='List Bullet')

try:
    doc.save(doc_path)
    print("Success")
except Exception as e:
    print(f"Error: {e}")
